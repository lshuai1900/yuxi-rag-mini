"""Test parsers: TXT, Markdown, DOCX, PDF, empty file error."""
import os
import tempfile
import pytest
from app.rag.parser.text_parser import TextParser
from app.rag.parser.markdown_parser import MarkdownParser
from app.rag.parser.docx_parser import DocxParser
from app.rag.parser.pdf_parser import PDFParser


@pytest.mark.asyncio
async def test_txt_parser():
    parser = TextParser()
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
        f.write("Hello world\nThis is a test file.")
        f.flush()
        result = await parser.parse(f.name, filename="test.txt")
    os.unlink(f.name)
    assert result.text == "Hello world\nThis is a test file."
    assert result.file_type == "txt"
    assert result.filename == "test.txt"


@pytest.mark.asyncio
async def test_txt_parser_utf8_sig():
    """Test TXT parser with UTF-8 BOM encoding."""
    parser = TextParser()
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="wb", delete=False) as f:
        f.write("\ufeffHello BOM world".encode("utf-8-sig"))
        f.flush()
        result = await parser.parse(f.name, filename="bom.txt")
    os.unlink(f.name)
    assert "Hello BOM world" in result.text


@pytest.mark.asyncio
async def test_txt_parser_gbk():
    """Test TXT parser with GBK encoding fallback."""
    parser = TextParser()
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="wb", delete=False) as f:
        f.write("中文GBK编码测试".encode("gbk"))
        f.flush()
        result = await parser.parse(f.name, filename="gbk.txt")
    os.unlink(f.name)
    assert "中文" in result.text


@pytest.mark.asyncio
async def test_markdown_parser():
    parser = MarkdownParser()
    content = "# Title\n\nParagraph text\n\n## Subtitle\n\nMore text"
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", encoding="utf-8", delete=False) as f:
        f.write(content)
        f.flush()
        result = await parser.parse(f.name, filename="test.md")
    os.unlink(f.name)
    assert result.text == content
    assert result.file_type == "markdown"
    assert "# Title" in result.text


@pytest.mark.asyncio
async def test_docx_parser():
    parser = DocxParser()
    from docx import Document
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc = Document()
        doc.add_heading("Test Heading", level=1)
        doc.add_paragraph("This is a test paragraph.")
        doc.save(f.name)
        result = await parser.parse(f.name, filename="test.docx")
    os.unlink(f.name)
    assert "Test Heading" in result.text
    assert "This is a test paragraph." in result.text
    assert result.file_type == "docx"


@pytest.mark.asyncio
async def test_pdf_parser_with_text():
    """Test PDF parser with a text-based PDF."""
    parser = PDFParser()
    # Create a simple PDF with reportlab if available, otherwise skip
    try:
        from reportlab.pdfgen import canvas
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            c = canvas.Canvas(f.name)
            c.drawString(100, 750, "Hello PDF World")
            c.save()
            result = await parser.parse(f.name, filename="test.pdf")
        os.unlink(f.name)
        assert "Hello PDF World" in result.text
        assert result.file_type == "pdf"
        assert len(result.pages) >= 1
        assert result.pages[0]["page_number"] == 1
    except ImportError:
        pytest.skip("reportlab not available for PDF test generation")


@pytest.mark.asyncio
async def test_empty_file_error():
    """Test that parsing an empty file raises a clear error."""
    parser = TextParser()
    with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", encoding="utf-8", delete=False) as f:
        f.write("")
        f.flush()
        with pytest.raises(ValueError, match="No extractable content"):
            await parser.parse(f.name, filename="empty.txt")
    os.unlink(f.name)


@pytest.mark.asyncio
async def test_empty_markdown_error():
    """Test that parsing an empty markdown file raises a clear error."""
    parser = MarkdownParser()
    with tempfile.NamedTemporaryFile(suffix=".md", mode="w", encoding="utf-8", delete=False) as f:
        f.write("   \n  \n  ")
        f.flush()
        with pytest.raises(ValueError, match="No extractable content"):
            await parser.parse(f.name, filename="empty.md")
    os.unlink(f.name)
