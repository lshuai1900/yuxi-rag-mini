import asyncio
from app.rag.parser.base import BaseParser, ParseResult


class DocxParser(BaseParser):
    async def parse(self, file_path: str, filename: str = "") -> ParseResult:
        return await asyncio.to_thread(self._parse_sync, file_path, filename)

    def _parse_sync(self, file_path: str, filename: str = "") -> ParseResult:
        from docx import Document
        document = Document(file_path)
        blocks = []

        # Process paragraphs in document order, skip empty ones
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # Preserve heading style info if available
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading ", "").replace("Heading", ""))
                    prefix = "#" * min(level, 6)
                    blocks.append(f"{prefix} {text}")
                except (ValueError, AttributeError):
                    blocks.append(text)
            else:
                blocks.append(text)

        # Process tables
        for table in document.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append(cells)
            if not rows:
                continue
            header = rows[0]
            blocks.append(f"| {' | '.join(header)} |")
            blocks.append(f"| {' | '.join(['---'] * len(header))} |")
            for row in rows[1:]:
                normalized = row + [""] * (len(header) - len(row))
                blocks.append(f"| {' | '.join(normalized[:len(header)])} |")
            blocks.append("")

        full_text = "\n\n".join(blocks).strip()

        if not full_text:
            raise ValueError(f"No extractable content found in {filename or 'DOCX file'}. The file may be empty.")

        return ParseResult(
            text=full_text,
            filename=filename,
            file_type="docx",
            pages=[],
            metadata={},
        )

    def supported_extensions(self) -> list[str]:
        return [".docx"]
